"""
DOCX export — exports research papers in Microsoft Word format.
"""

from __future__ import annotations

import logging
from io import BytesIO
from typing import Optional, List, Dict, Any, Union
from datetime import datetime

from backend.export.base import Exporter

logger = logging.getLogger(__name__)


class DOCXExporter(Exporter):
    """
    Exports research papers in Microsoft Word DOCX format.
    Note: This is a simplified version - a real implementation would use python-docx library.
    """

    def __init__(self) -> None:
        super().__init__("docx")

    def export(
        self,
        paper_data: Dict[str, Any],
        output_path: Optional[str] = None,
        include_bibliography: bool = True,
        **kwargs
    ) -> Union[str, bytes]:
        """
        Export research paper data to a DOCX-compatible format (simulated).
        
        Args:
            paper_data: Dictionary containing paper data (sections, citations, etc.)
            output_path: Optional path to save the file
            include_bibliography: Whether to include bibliography section
            **kwargs: Additional format-specific parameters
            
        Returns:
            DOCX content representation as string (simulated)
        """
        try:
            return self._export_real_docx(paper_data, output_path, include_bibliography)
        except Exception as exc:
            logger.warning("Falling back to textual DOCX representation: %s", exc)
            return self._export_textual_representation(paper_data, output_path, include_bibliography)

    def _export_real_docx(
        self,
        paper_data: Dict[str, Any],
        output_path: Optional[str],
        include_bibliography: bool,
    ) -> bytes:
        """Create a real DOCX document using python-docx when available."""
        from docx import Document

        document = Document()
        title = paper_data.get("title", "Untitled Research Paper")
        authors = paper_data.get("authors", [])
        abstract = paper_data.get("abstract", "")

        document.add_heading(title, level=0)
        if authors:
            document.add_paragraph(", ".join(authors), style=None)
        if abstract:
            document.add_heading("Abstract", level=1)
            document.add_paragraph(abstract)

        for section in paper_data.get("sections", []) or []:
            document.add_heading(section.get("title", "Untitled Section"), level=1)
            for paragraph in str(section.get("content", "")).split("\n\n"):
                if paragraph.strip():
                    document.add_paragraph(paragraph.strip())

        citations = paper_data.get("citations", []) or []
        if include_bibliography and citations:
            document.add_heading("References", level=1)
            for index, citation in enumerate(citations, 1):
                document.add_paragraph(f"{index}. {self._format_citation_for_docx(citation)}")

        document.add_paragraph(f"Exported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        buffer = BytesIO()
        document.save(buffer)
        content = buffer.getvalue()

        if output_path:
            with open(output_path, "wb") as handle:
                handle.write(content)
            logger.info("Saved DOCX export to %s", output_path)

        return content

    def _export_textual_representation(
        self,
        paper_data: Dict[str, Any],
        output_path: Optional[str],
        include_bibliography: bool,
    ) -> str:
        """Fallback structured text representation for environments without DOCX support."""
        content_parts = []
        
        # Title and metadata
        title = paper_data.get("title", "Untitled Research Paper")
        authors = paper_data.get("authors", [])
        abstract = paper_data.get("abstract", "")
        
        content_parts.append("DOCUMENT STRUCTURE")
        content_parts.append("=" * 50)
        content_parts.append(f"Title: {title}")
        content_parts.append("")
        
        if authors:
            authors_str = ", ".join(authors)
            content_parts.append(f"Authors: {authors_str}")
            content_parts.append("")
            
        if abstract:
            content_parts.append("ABSTRACT")
            content_parts.append("-" * 20)
            content_parts.append(abstract)
            content_parts.append("")
        
        # Sections
        sections = paper_data.get("sections", [])
        for i, section in enumerate(sections, 1):
            section_title = section.get("title", "Untitled Section")
            section_content = section.get("content", "")
            
            content_parts.append(f"SECTION {i}: {section_title}")
            content_parts.append("-" * (30 + len(str(i))))
            content_parts.append(section_content)
            content_parts.append("")
            
        # Bibliography
        if include_bibliography:
            citations = paper_data.get("citations", [])
            if citations:
                content_parts.append("REFERENCES")
                content_parts.append("=" * 50)
                
                for i, citation in enumerate(citations, 1):
                    citation_str = self._format_citation_for_docx(citation)
                    content_parts.append(f"{i}. {citation_str}")
                content_parts.append("")
        
        # Add timestamp
        content_parts.append("=" * 50)
        content_parts.append(f"Exported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content_parts.append("Note: This is a structural representation.")
        content_parts.append("A real DOCX export would use python-docx library.")
        
        docx_content = "\n".join(content_parts)
        
        # Save if path provided
        if output_path:
            try:
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(docx_content)
                logger.info(f"Saved DOCX representation to {output_path}")
            except Exception as e:
                logger.error(f"Failed to save DOCX representation: {e}")
                raise
        
        return docx_content

    def _format_citation_for_docx(self, citation: Dict[str, Any]) -> str:
        """
        Format a citation in DOCX-friendly format.
        
        Args:
            citation: Citation dictionary
            
        Returns:
            Formatted citation string
        """
        title = citation.get("title", "Unknown")
        authors = citation.get("authors", [])
        year = citation.get("year", "Unknown")
        
        authors_str = ", ".join(authors) if authors else "Unknown Author"
        return f"{authors_str} ({year}). {title}."

    def get_supported_formats(self) -> List[str]:
        """Get list of supported export formats for this exporter."""
        return ["docx"]

    def get_file_extension(self) -> str:
        """Get file extension for this exporter."""
        return ".docx"


# Global instance
docx_exporter = DOCXExporter()